from pathlib import Path
from docx import Document

template = Path(r"C:\Users\CasaT\Downloads\Quaerens_Master_Assessment_Template_v2.docx")
doc = Document(str(template))

print("PARAGRAPHS")
for i, p in enumerate(doc.paragraphs[:80], 1):
    text = " ".join(p.text.split())
    if text:
        print(f"{i:03d} | {p.style.name} | {text[:180]}")

print("\nTABLES")
for ti, table in enumerate(doc.tables, 1):
    print(f"Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows[:5], 1):
        vals = [" ".join(c.text.split())[:80] for c in row.cells]
        print(f"  {ri}: {vals}")
