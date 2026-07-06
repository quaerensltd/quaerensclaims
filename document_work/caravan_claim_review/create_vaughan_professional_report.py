from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\CasaT\quaerensclaims\document_work\caravan_claim_review\01_Kennedy_Vaughan_Haven_Caravan_Assessment_Report_Professional.docx")

BLUE = "005BAA"
DARK = "0B1F3A"
LIGHT_BLUE = "EAF4FF"
PALE = "F7FBFF"
GREY = "F2F4F7"
GOLD = "FFF2CC"
GREEN = "0B7A3B"
RED = "9B1C1C"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color="D9E5F2", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        elem = tc_borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            tc_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def cell_text(cell, text, bold=False, size=10, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    borders(cell)
    return p


def para(doc, text="", style=None, bold_start=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else DARK)
        r.font.size = Pt(15 if level == 1 else 12)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    p.add_run(text)
    return p


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    p.add_run(text)
    return p


def simple_table(doc, headers, rows, widths=None, header_fill=GREY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell_text(cell, h, bold=True, size=9)
        shade(cell, header_fill)
        if widths:
            cell.width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell_text(cells[i], val, size=9)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def note_box(doc, text, title=None, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    borders(cell, color="BBD7F2", size="10")
    if title:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(BLUE)
        p2 = cell.add_paragraph()
    else:
        p2 = cell.paragraphs[0]
    p2.paragraph_format.space_after = Pt(3)
    p2.paragraph_format.line_spacing = 1.1
    r = p2.add_run(text)
    r.font.size = Pt(10)
    doc.add_paragraph()


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.75)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(7)

# Title band
title_tbl = doc.add_table(rows=1, cols=1)
title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = title_tbl.cell(0, 0)
shade(cell, DARK)
borders(cell, color=DARK, size="12")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("QUAERENS LTD\nProfessional Client Assessment Report\nStatic Caravan / Holiday Park Claim Review")
r.bold = True
r.font.color.rgb = RGBColor.from_string("FFFFFF")
r.font.size = Pt(16)
doc.add_paragraph()

top = doc.add_table(rows=1, cols=2)
top.alignment = WD_TABLE_ALIGNMENT.CENTER
top.autofit = False
top.cell(0, 0).width = Inches(3.1)
top.cell(0, 1).width = Inches(3.7)
shade(top.cell(0, 0), PALE)
shade(top.cell(0, 1), PALE)
borders(top.cell(0, 0))
borders(top.cell(0, 1))
cell_text(top.cell(0, 0), "Client\nKennedy Vaughan\nAddress / contact details: TBC from client file", bold=False, size=10)
cell_text(top.cell(0, 1), "Prepared for\nInitial claim assessment and evidence review\nDate prepared: 3 July 2026", bold=False, size=10)
doc.add_paragraph()

note_box(
    doc,
    "This assessment is preliminary and based on the litigation bundle, SAR material and registers provided. It is not legal advice and does not guarantee recovery. The case should be referred to an appropriate legal partner once the evidence index, chronology and Schedule of Loss are complete.",
    title="Important status note:",
    fill=GOLD,
)

heading(doc, "Client Information Matrix")
simple_table(
    doc,
    ["Field", "Information"],
    [
        ("Client Name", "Kennedy Vaughan"),
        ("Address", "TBC - not provided in the reviewed bundle"),
        ("Telephone", "TBC"),
        ("Email", "TBC"),
        ("Product", "Static holiday caravan / replacement caravan"),
        ("Supplier / Park Operator", "Haven Holidays Limited / Bourne Leisure Ltd - exact contracting entity to be checked"),
        ("Park / Location", "Haven / Hopton referenced in additional bundle; precise park and pitch details TBC"),
        ("Relevant Period", "First purchase around 2016/2017; upgrade date and handover date TBC"),
        ("Documents Reviewed", "Haven litigation bundle sections 1-20, master chronology, SAR register, SAR analysis and authorities register"),
        ("Primary Complaint", "Alleged pre-contract assurances that defects would be rectified before handover, followed by unresolved defects and prolonged complaint/SAR history"),
        ("Claim Status", "Detailed narrative prepared; primary exhibits, figures and final court bundle references still require completion"),
    ],
    [Inches(1.65), Inches(5.05)],
)

heading(doc, "Executive Summary")
para(doc, "Mr Kennedy Vaughan appears to have a potentially arguable static caravan claim against Haven Holidays Limited / Bourne Leisure Ltd arising from the purchase of a replacement caravan. The central allegation is that Haven's representative induced the purchase by giving clear assurances that identified defects would be rectified before handover. The claimant says those assurances were relied upon and were not honoured.")
para(doc, "The additional material strengthens the evidence-management side of the case. It identifies a Subject Access Request history, delayed disclosure, internal Haven records and a legal authorities register covering the Consumer Rights Act 2015, Misrepresentation Act 1967, Consumer Protection from Unfair Trading Regulations 2008, statutory interest and civil procedure.")
para(doc, "The matter is not yet ready to be presented as a fully quantified legal claim. The report should be treated as a detailed preliminary assessment. The strongest next step is to convert the existing narrative bundle into a solicitor-ready evidence pack with exhibit numbers, exact dates, signed witness evidence and a properly evidenced Schedule of Loss.")

note_box(
    doc,
    "Preliminary claim strength: potentially moderate to strong, subject to evidence. The strongest points are the alleged pre-handover assurances, pre-contract knowledge of defects, reliance, continuing unresolved defects, and Haven's own internal/SAR records. The main risks are limitation, missing financial figures, and the need to prove the exact words or written record of the assurances.",
    title="Preliminary claim strength:",
    fill=LIGHT_BLUE,
)

heading(doc, "Background & Chronology")
simple_table(
    doc,
    ["Date / Stage", "Event / Relevance"],
    [
        ("2015", "Bundle states Mr Vaughan suffered three brain aneurysms, with long-term effects including memory impairment."),
        ("2016 / 2017", "Family purchased first Haven caravan for private family use, during a period affected by their son's lymphoma treatment."),
        ("Relationship with Haven", "The bundle states that a relationship of trust developed with Haven's representative, Aaron, who knew the family's personal circumstances."),
        ("Upgrade discussions", "Haven allegedly encouraged an upgrade to a replacement caravan."),
        ("Pre-signature inspection", "Mr and Mrs Vaughan allegedly inspected the replacement caravan and identified numerous defects before signing."),
        ("Assurances before purchase", "The claimant says Aaron assured them all identified defects / snagging would be rectified before occupation or handover."),
        ("Handover / occupation", "The claimant says the caravan was not completed as promised and numerous defects remained."),
        ("Complaint period", "The claimant made repeated complaints and gave Haven opportunities to remedy the issues over an extended period."),
        ("June 2025 TBC", "SAR register says an initial Subject Access Request was submitted to Haven / Bourne Leisure; exact date to confirm."),
        ("17 January 2026", "SAR register says disclosure was received after significant delay; to check against the source email/bundle."),
        ("20 March 2026", "SAR register refers to Haven correspondence concerning delay/admin handling of the SAR; to exhibit."),
        ("Current", "Litigation bundle and evidence analysis prepared; final exhibits, page references and loss figures remain to be completed."),
    ],
    [Inches(1.55), Inches(5.15)],
)

heading(doc, "Technical Findings Analysis")
para(doc, "The technical issue is the condition of the replacement caravan before and after purchase. The available bundle repeatedly states that numerous defects were identified before signature, that Haven was made aware of them, and that completion of remedial works was promised before handover.")
heading(doc, "Defect and snagging concerns", 2)
for item in [
    "The bundle refers to approximately seventy-one defects, but the final defect schedule still needs to be attached and cross-referenced.",
    "The most important proof will be a pre-handover snagging list, photographs, emails and Haven repair/work-order records.",
    "Each defect should be mapped to: date identified, whether it was pre-contract, who was told, what was promised, what was completed and what remained outstanding.",
    "If independent inspection evidence exists, it should be prioritised because it may reduce reliance on recollection alone.",
]:
    bullet(doc, item)
heading(doc, "Repair and remedial work concerns", 2)
para(doc, "The claimant's account is that Haven was given repeated opportunities to resolve the problem. The final evidence pack should distinguish between defects that were never repaired, defects repaired late, defects repaired poorly, and defects that caused continuing loss of use or enjoyment.")

note_box(
    doc,
    "Technical evidence priority: the case needs a single master defect schedule supported by dated photographs, emails, snagging records, repair logs and Haven internal notes. Without that schedule, the legal narrative remains much harder to prove.",
    title="Technical evidence priority:",
    fill=LIGHT_BLUE,
)

heading(doc, "Legal / Contract / Consumer Analysis")
para(doc, "This is not a finance-only complaint. The principal routes are likely to be contractual, consumer and representation-based. The correct route should be confirmed by a solicitor after reviewing the purchase agreement, written terms, warranty, park rules and any disclaimer or inspection wording.")
for item in [
    "Misrepresentation / inducement: potentially strong if clear pre-contract assurances can be proved and the claimant relied on them when purchasing.",
    "Breach of contract: potentially arguable if the promise to complete remedial works formed part of the agreed transaction.",
    "Consumer Rights Act 2015: relevant to satisfactory quality, fitness, description and pre-contract information where a trader supplied goods/services to a consumer.",
    "Consumer Protection from Unfair Trading Regulations 2008: potentially relevant if there were misleading actions, omissions or pressure tactics during the sales process.",
    "Misrepresentation Act 1967: listed in the client's authorities register and relevant to alleged false or unfulfilled representations.",
    "Data protection / SAR issues: useful as supporting evidence and potentially a separate complaint route, but not the primary caravan sale cause of action.",
]:
    bullet(doc, item)

heading(doc, "SAR / Disclosure / Internal Records Analysis")
para(doc, "The additional bundle places greater emphasis on the Subject Access Request and Haven's internal records. This is important because the claimant's case should not rely only on memory. Haven's own records may confirm dates, conversations, complaint handling, repair history and internal knowledge.")
simple_table(
    doc,
    ["SAR Ref", "Point from new material", "Assessment"],
    [
        ("SAR001", "Initial Subject Access Request submitted, June 2025 TBC.", "Confirm exact date and attach original request."),
        ("SAR002", "Acknowledgement / delay issue during 2025.", "Attach Haven acknowledgement and chase emails."),
        ("SAR003", "Disclosure said to have been received 17 January 2026.", "Index received documents and identify strongest internal records."),
        ("SAR004", "Haven apology / admin error correspondence dated 20 March 2026.", "Potentially useful conduct evidence; attach and summarise."),
        ("SAR005", "Missing / incomplete disclosure review pending.", "Prepare missing-document schedule if relevant."),
        ("SAR006", "Internal Haven records received through SAR disclosure.", "Compare against claimant chronology and exhibit contradictions/support."),
        ("SAR007", "SAR relevance to claim.", "Use to support internal knowledge, complaint handling and credibility, not as a substitute for loss evidence."),
    ],
    [Inches(0.9), Inches(2.4), Inches(3.4)],
)

heading(doc, "Financial Exposure")
para(doc, "Financial exposure remains TBC because the extracted bundle does not provide a complete purchase price, site-fee history, insurance position, repair costs, removal/disposal figures or residual value evidence. Any report to solicitors should include a fully evidenced Schedule of Loss.")
simple_table(
    doc,
    ["Field", "Amount", "Basis / Comment"],
    [
        ("Replacement caravan purchase cost", "TBC", "Purchase agreement, invoice and payment evidence required."),
        ("Part-exchange / trade-in value", "TBC", "Needed to understand net purchase loss."),
        ("Annual site / pitch fees", "TBC", "Statements and dates required; only claim where linked to loss period and legally recoverable."),
        ("Insurance / ownership costs", "TBC", "Policy documents and payment records required."),
        ("Repair / inspection costs", "TBC", "Invoices, reports and receipts required."),
        ("Removal / disposal / residual value", "TBC", "Bundle says caravan was ultimately removed; full documentation required."),
        ("Loss of use / enjoyment", "TBC", "May be relevant, but recoverability and valuation require legal advice."),
        ("Current quantified exposure", "Not yet reliable", "Narrative exists, but the financial schedule is incomplete."),
    ],
    [Inches(1.7), Inches(1.2), Inches(3.8)],
)

heading(doc, "Valuation Matrix")
para(doc, "The following matrix is indicative only. It should not be presented as a guaranteed outcome. Final valuation depends on contract terms, limitation, liability evidence and documented loss.")
simple_table(
    doc,
    ["Scenario", "Indicative Position", "Rationale"],
    [
        ("Conservative", "Documented direct outlay only", "Recovery limited to clearly evidenced costs directly linked to the unresolved defect issue."),
        ("Moderate", "Direct outlay plus consequential costs", "Applies if purchase documents, repair history and complaint records support breach/misrepresentation."),
        ("Strong", "Material recovery claim", "Applies if Haven's own records support pre-contract assurances, unresolved defects and loss causation."),
        ("Exceptional", "Only after solicitor review", "Requires clear liability, strong documents, manageable limitation and a complete Schedule of Loss."),
    ],
    [Inches(1.25), Inches(1.75), Inches(3.7)],
)

heading(doc, "Evidence Required")
para(doc, "The claim should now move from narrative bundle to exhibit-backed case pack. The objective is to prove what was promised, what was supplied, what remained defective and what each loss figure represents.")
for item in [
    "Signed purchase agreement, terms and conditions, warranty, park rules and any disclaimer/inspection wording.",
    "Invoice, payment records, part-exchange documents and any finance paperwork.",
    "Pre-handover snagging list and any written notes showing defects identified before signature.",
    "Photographs/videos of each defect, preferably dated and matched to the defect schedule.",
    "Emails, letters or messages recording Aaron/Haven's assurances.",
    "Haven work orders, repair logs, inspection notes and internal records disclosed through SAR.",
    "Full SAR correspondence: request, acknowledgement, chasers, disclosure, apology/admin error correspondence and missing disclosure review.",
    "Formal complaint letters and Haven responses.",
    "Annual site fee, insurance, repair, removal/disposal and residual value evidence.",
    "Signed witness statements from Mr Vaughan and Mrs Vaughan.",
]:
    number(doc, item)

heading(doc, "Preliminary Assessment")
simple_table(
    doc,
    ["Assessment Factor", "Preliminary Position"],
    [
        ("Misrepresentation / Reliance", "Potentially strong if the assurances are supported by witness statements, sales records, messages or Haven internal notes."),
        ("Contract / CRA Position", "Potentially arguable, but the purchase agreement and warranty wording must be reviewed."),
        ("Technical Evidence", "Promising narrative, but requires a defect-by-defect exhibit schedule."),
        ("SAR / Disclosure Evidence", "Now stronger. Additional bundle provides a SAR register, disclosure history and internal-records analysis framework."),
        ("Loss Evidence", "Currently weak to moderate because figures remain TBC."),
        ("Limitation Risk", "High until exact purchase, handover and complaint dates are reviewed by solicitors."),
        ("Overall View", "Good candidate for legal partner review once the evidence and loss schedules are completed."),
    ],
    [Inches(2.0), Inches(4.7)],
)

heading(doc, "Recommendations & Next Steps")
for item in [
    "Create a single chronological case summary with exact dates and exhibit references.",
    "Prepare a master defect schedule covering all alleged defects and repair history.",
    "Complete the SAR evidence schedule using references KV-SAR-001 onwards.",
    "Build a fully evidenced Schedule of Loss before any settlement figure is discussed.",
    "Confirm the legal entity: Haven Holidays Limited, Bourne Leisure Ltd, or another contracting party shown on the sales paperwork.",
    "Send the completed pack to a legal partner for limitation, liability and quantum advice.",
]:
    number(doc, item)

note_box(
    doc,
    "Recommended case direction: proceed to evidence indexing and solicitor review. The case has a detailed narrative and useful SAR/internal-records angle, but the professional strength of the file will depend on matching each allegation to exhibits and each loss to proof.",
    title="Recommended case direction:",
    fill=LIGHT_BLUE,
)

heading(doc, "Specialised Caravan / Holiday Park Assessment Points")
for item in [
    "Sales representation analysis should focus on exactly what was said before signature and whether those statements induced the purchase.",
    "Condition analysis should separate pre-contract defects, handover defects and later developing issues.",
    "Holiday park ownership analysis should include site fees, pitch/licence terms, park rules, warranties and any restriction on remedies.",
    "SAR analysis should be used to compare Haven's internal records against the claimant's chronology.",
    "Loss analysis should avoid unsupported figures and should separate purchase loss, ongoing ownership costs, repair/removal costs and loss of use.",
]:
    bullet(doc, item)

end = doc.add_paragraph("End of Preliminary Assessment Report")
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in end.runs:
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Quaerens Ltd - Professional Client Assessment Report")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor.from_string("666666")

doc.save(OUT)
print(OUT)
