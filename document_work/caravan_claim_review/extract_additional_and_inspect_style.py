from pathlib import Path
from docx import Document

BASE = Path(r"C:\Users\CasaT\quaerensclaims\document_work\caravan_claim_review")
ADDITIONAL = BASE / "additional_bundle"
STYLE_DOC = Path(r"C:\Users\CasaT\Desktop\jeroen\temp\Irene_Denison_Quaerens_Assessment_Report_Professional.docx")
OUT = BASE / "additional_bundle_text_digest.txt"
STYLE_OUT = BASE / "irene_style_inspection.txt"

def docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        t = " ".join(p.text.split())
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            vals = [" ".join(c.text.split()) for c in row.cells]
            line = " | ".join(v for v in vals if v)
            if line:
                parts.append(line)
    return "\n".join(parts)

with OUT.open("w", encoding="utf-8") as f:
    for path in sorted(ADDITIONAL.glob("*.docx")):
        f.write("\n" + "=" * 90 + "\n")
        f.write(path.name + "\n")
        f.write("=" * 90 + "\n")
        f.write(docx_text(path))
        f.write("\n")

style_doc = Document(str(STYLE_DOC))
with STYLE_OUT.open("w", encoding="utf-8") as f:
    f.write("PARAGRAPHS\n")
    for i, p in enumerate(style_doc.paragraphs[:140], 1):
        text = " ".join(p.text.split())
        if text:
            f.write(f"{i:03d} | {p.style.name} | {text[:220]}\n")
    f.write("\nTABLES\n")
    for ti, table in enumerate(style_doc.tables, 1):
        f.write(f"Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols\n")
        for ri, row in enumerate(table.rows[:8], 1):
            vals = [" ".join(c.text.split())[:100] for c in row.cells]
            f.write(f"  {ri}: {vals}\n")

print(OUT)
print(STYLE_OUT)
