from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEMPLATE = Path(r"C:\Users\CasaT\Downloads\Quaerens_Master_Assessment_Template_v2.docx")
OUT = Path(r"C:\Users\CasaT\quaerensclaims\document_work\caravan_claim_review\01_Kennedy_Vaughan_Haven_Caravan_Assessment_Report.docx")


def set_para_text(paragraph, text, bold=False):
    paragraph.text = ""
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


def insert_after(paragraph, text="", style=None, bold_start=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    if style:
        p.style = style
    p.paragraph_format.space_after = Pt(6)
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def remove_para(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def add_items_after(anchor, items):
    p = anchor
    for item in items:
        p = insert_after(p, item, style="List Bullet")
    return p


def fill_table(table, rows):
    for r_idx, row in enumerate(rows):
        if r_idx >= len(table.rows):
            table.add_row()
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            cells[c_idx].text = str(value)
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)


doc = Document(str(TEMPLATE))

# Remove the generic variants section from the master template.
for p in list(doc.paragraphs):
    if "SPECIALISED VARIANTS" in p.text or p.text.startswith("A. Spray Foam Template"):
        remove_para(p)

paras = doc.paragraphs
set_para_text(paras[0], "QUAERENS LTD CARAVAN CLAIM ASSESSMENT REPORT", bold=True)
paras[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_text(paras[1], "Initial Assessment - Kennedy Vaughan v Haven Holidays Limited / Bourne Leisure Ltd")
paras[1].alignment = WD_ALIGN_PARAGRAPH.CENTER

fill_table(
    doc.tables[0],
    [
        ("Field", "Information"),
        ("Client Name", "Kennedy Vaughan"),
        ("Address", "Not provided in the reviewed bundle"),
        ("Telephone", "Not provided in the reviewed bundle"),
        ("Email", "Not provided in the reviewed bundle"),
        ("Respondent / Trader", "Haven Holidays Limited / Bourne Leisure Ltd - exact contracting entity TBC"),
        ("Product / Issue", "Static holiday caravan purchase / replacement caravan defects and alleged pre-contract assurances"),
        ("Purchase / Event Date", "Exact dates TBC; bundle refers to first purchase around 2016/2017 and later upgrade discussions"),
        ("Current Status", "Draft litigation bundle reviewed; exhibits, dates and Schedule of Loss require completion"),
        ("Assessment Date", "3 July 2026"),
    ],
)

fill_table(
    doc.tables[1],
    [
        ("Field", "Information"),
        ("Purchase Cost", "TBC - purchase agreement/invoice required"),
        ("Site Fees / Pitch Fees", "TBC - annual statements required"),
        ("Insurance / Ownership Costs", "TBC - policy and payment records required"),
        ("Repair / Inspection Costs", "TBC - invoices and receipts required"),
        ("Removal / Disposal / Residual Value", "TBC - caravan was reportedly removed; full evidence required"),
        ("Complaint / Document Costs", "TBC - recoverability requires legal review"),
        ("Overall Financial Exposure", "Potentially material, but not yet quantifiable from the extracted bundle"),
    ],
)

fill_table(
    doc.tables[2],
    [
        ("Field", "Information"),
        ("Conservative", "Evidence supports a long complaint history, but loss figures and contract terms remain TBC."),
        ("Moderate", "If documents prove the pre-contract defect list and clear assurances, a viable consumer claim may exist."),
        ("Strong", "If Haven's own records, emails or snagging documents confirm the assurances and unresolved defects."),
        ("Exceptional", "Only if liability evidence is clear, losses are fully evidenced, limitation is manageable and legal review supports proceedings."),
    ],
)

content = {
    "EXECUTIVE SUMMARY": [
        "Based on the bundle reviewed, this appears to be a potentially arguable caravan claim suitable for further legal partner review. The strongest point is not simply that the caravan had defects; it is that Mr and Mrs Vaughan allegedly identified those defects before purchase, Haven's representative allegedly promised that all identified remedial works would be completed before handover, and the claimant says he relied on those promises when proceeding with the replacement caravan purchase.",
        "The documents reviewed are mostly a drafted litigation bundle rather than a fully indexed evidence pack. The narrative is coherent and repeats the same central allegation across the sections reviewed, but exact dates, purchase figures, exhibit references and the Schedule of Loss are still incomplete. Those gaps need to be closed before any formal letter before action, settlement proposal or referral to solicitors.",
    ],
    "BACKGROUND & CHRONOLOGY": [
        "The claimant is Kennedy Vaughan. The proposed defendant is identified in the bundle as Haven Holidays Limited / Bourne Leisure Ltd. The claimant and his wife purchased a static caravan for private family use, not as an investment or commercial letting business.",
        "The bundle explains that the family's circumstances were sensitive. It refers to their son's lymphoma treatment and Mr Vaughan's own history of three brain aneurysms and memory issues. The documents state that Haven's representative knew these circumstances and that a relationship of trust developed during the sales relationship.",
        "After an initial caravan purchase, Haven allegedly encouraged an upgrade to a replacement caravan. Before signature, Mr and Mrs Vaughan allegedly inspected the replacement caravan, identified numerous defects and discussed them with Haven's representative, Aaron. The claimant says Aaron assured them that every defect would be rectified before handover.",
        "The claimant says the purchase proceeded because those assurances were accepted. After handover, the promised condition was allegedly not achieved. The bundle refers to approximately seventy-one defects, repeated repair requests, formal complaints, Subject Access Request disclosure, eventual loss of confidence and preparation for litigation.",
    ],
    "TECHNICAL FINDINGS ANALYSIS": [
        "The technical evidence is described rather than fully exhibited in the extracted bundle. The bundle refers to snagging records, photographs, inspection material, complaint correspondence and Haven internal records. Those materials will be essential because the claim depends on proving what defects existed before purchase, what Haven knew, what was promised and what remained unresolved after handover.",
        "The most important technical task is to create a clear defect schedule. This should show each defect, whether it was identified before signature, whether Haven accepted it or promised repair, the promised completion date if any, what actually happened, and what evidence supports each point.",
    ],
    "MORTGAGE / LENDING / PROPERTY VALUE ANALYSIS": [
        "This is not a mortgage or lending claim on the documents reviewed. For this caravan matter, the equivalent analysis is contract, consumer rights and representation analysis.",
        "The possible routes identified from the bundle are: misrepresentation and inducement to contract, breach of contract/failure to perform agreed remedial works, potential Consumer Rights Act 2015 issues relating to satisfactory quality and pre-contract information, and complaint handling/SAR delay as supporting background evidence.",
        "The correct legal entity and contractual terms must be checked from the purchase agreement. Any warranty, exclusion clause, 'sold as seen' wording, inspection clause or complaint time limit may materially affect the analysis.",
    ],
    "FINANCIAL EXPOSURE": [
        "Financial exposure cannot yet be reliably calculated from the extracted bundle. The bundle says a Schedule of Loss is still to be completed. Loss evidence should include the purchase price, any part-exchange value, finance records if relevant, annual site fees, insurance, repair costs, inspection costs, removal/disposal costs, residual value and any loss of use or enjoyment claimed.",
        "The claim should avoid broad unsupported loss figures. Each claimed figure should be supported by an invoice, statement, receipt, contract, valuation, bank record or other documentary proof.",
    ],
    "VALUATION MATRIX": [
        "The valuation position is currently TBC. For a caravan claim, the key valuation questions are: what was paid, what the caravan was represented to be worth in the promised condition, what it was actually worth given the outstanding defects, what continuing costs were incurred, and what value was recovered when the caravan was removed or disposed of.",
    ],
    "EVIDENCE REQUIRED": [
        "The following documents should be obtained and indexed before the matter is presented as litigation-ready:",
    ],
    "PRELIMINARY ASSESSMENT": [
        "Preliminary merits: moderate to potentially strong, subject to evidence. The case narrative is consistent and potentially compelling, especially if documentary records confirm that defects were identified before purchase and that Haven gave clear assurances before signature.",
        "Main risk: evidence readiness and limitation. The bundle is historic and repeatedly states that dates and exhibits are still to be finalised. A solicitor should review limitation before any formal action is taken.",
        "Operational view: this is suitable for a legal partner referral once the evidence schedule and loss schedule are complete. Quaerens should not present the claim as guaranteed or fully quantified at this stage.",
    ],
    "RECOMMENDATIONS & NEXT STEPS": [
        "Recommended action: complete the evidence pack, quantify the loss, and submit the matter for solicitor review before any formal letter before action is issued.",
    ],
}

evidence_items = [
    "Signed purchase agreement, terms and conditions, warranty and any 'sold as seen' or inspection wording.",
    "Invoice, finance agreement, part-exchange evidence and payment records.",
    "Pre-handover snagging list or notes showing defects identified before signature.",
    "Photographs/video evidence of defects, ideally dated and matched to the defect schedule.",
    "Emails, WhatsApp messages or notes recording Aaron/Haven's assurances.",
    "Haven complaint responses, repair logs, work orders and internal records from SAR disclosure.",
    "Evidence of annual site fees, insurance, repair costs and removal/disposal/residual value.",
    "Witness statements from Mr Vaughan and Mrs Vaughan covering inspection, promises, reliance and impact.",
]

next_steps = [
    "Confirm the exact contracting entity and purchase date from the sale agreement.",
    "Build a dated chronology from first contact through upgrade, inspection, handover, complaints, SAR, removal and solicitor involvement.",
    "Prepare a defect-by-defect evidence schedule with exhibit numbers.",
    "Prepare a quantified Schedule of Loss with documentary support for every figure.",
    "Obtain signed witness statements from Mr and Mrs Vaughan.",
    "Refer the completed pack to a legal partner for limitation and prospects advice.",
]

for heading, texts in content.items():
    idx = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == heading), None)
    if idx is None:
        continue
    anchor = doc.paragraphs[idx]
    # Remove the generic placeholder immediately after the heading where present.
    if idx + 1 < len(doc.paragraphs) and "[Insert detailed narrative here]" in doc.paragraphs[idx + 1].text:
        remove_para(doc.paragraphs[idx + 1])
    cur = anchor
    for text in texts:
        cur = insert_after(cur, text)
    if heading == "EVIDENCE REQUIRED":
        cur = add_items_after(cur, evidence_items)
    if heading == "RECOMMENDATIONS & NEXT STEPS":
        cur = add_items_after(cur, next_steps)
        cur = insert_after(
            cur,
            "Important note: This is an initial assessment only. Quaerens should use legal partners to proceed with suitable caravan cases where, after review, the evidence indicates that the client may have a viable claim.",
            bold_start="Important note:",
        )

doc.save(OUT)
print(OUT)
