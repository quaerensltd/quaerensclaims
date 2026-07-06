from __future__ import annotations

import json
import re
from pathlib import Path

import pdfplumber
from docx import Document


WORK = Path(r"C:\Users\CasaT\quaerensclaims\document_work\caravan_claim_review")
OUT = WORK / "extra_hopton_docs_digest.txt"
SUMMARY = WORK / "extra_hopton_docs_findings.json"

FILES = [
    Path(r"C:\Users\CasaT\Downloads\hopton .pdf"),
    Path(r"C:\Users\CasaT\Downloads\sar hopton"),
    Path(r"C:\Users\CasaT\Downloads\INSURANCE FOR VAN [56].pdf"),
    Path(r"C:\Users\CasaT\Downloads\complaint to hopton  .docx"),
]


def detect(path: Path) -> str:
    sig = path.read_bytes()[:8]
    if sig.startswith(b"%PDF"):
        return "pdf"
    if sig.startswith(b"PK"):
        return "docx_or_zip"
    return sig.hex()


def read_pdf(path: Path) -> str:
    parts = []
    try:
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(f"[Page {i}]\n{text.strip()}")
    except Exception as exc:
        return f"[UNREADABLE PDF: {exc}]"
    return "\n\n".join(parts)


def read_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
    except Exception as exc:
        return f"[UNREADABLE DOCX: {exc}]"
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip().replace("\n", " | ") for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def context_hits(text: str, patterns: list[str], radius: int = 240) -> list[str]:
    hits = []
    for pat in patterns:
        for match in re.finditer(pat, text, flags=re.IGNORECASE):
            start = max(0, match.start() - radius)
            end = min(len(text), match.end() + radius)
            snippet = re.sub(r"\s+", " ", text[start:end]).strip()
            if snippet not in hits:
                hits.append(snippet)
            if len(hits) >= 12:
                return hits
    return hits


def main() -> None:
    chunks = []
    file_info = []
    for file in FILES:
        if not file.exists():
            file_info.append({"path": str(file), "exists": False})
            continue
        kind = detect(file)
        if kind == "pdf":
            text = read_pdf(file)
        elif kind == "docx_or_zip":
            text = read_docx(file)
        else:
            text = file.read_text(encoding="utf-8", errors="ignore")
        file_info.append({"path": str(file), "exists": True, "kind": kind, "chars": len(text)})
        chunks.append(f"\n\n===== {file} ({kind}) =====\n{text}")

    all_text = "\n".join(chunks)
    OUT.write_text(all_text, encoding="utf-8", errors="ignore")

    categories = {
        "contact_identity": [r"\bKennedy\b", r"\bVaughan\b", r"\bNicola\b", r"\bemail\b", r"\bphone\b", r"\baddress\b"],
        "park_pitch_caravan": [r"\bHopton\b", r"\bpitch\b", r"\bplot\b", r"\bcaravan\b", r"\bmodel\b", r"\bserial\b", r"\bCRiS\b", r"\bABI\b", r"\bSwift\b", r"\bWillerby\b"],
        "insurance": [r"\binsurance\b", r"\bpolicy\b", r"\bpremium\b", r"\binsurer\b", r"\bcover\b"],
        "prices_losses": [r"£\s?\d", r"\bprice\b", r"\bdeposit\b", r"\brefund\b", r"\bloss\b", r"\bcompensation\b", r"\bdamages\b"],
        "complaint_sar": [r"\bcomplaint\b", r"\bsubject access\b", r"\bSAR\b", r"\bGDPR\b", r"\bICO\b", r"\bprivacy\b"],
        "defects_promises": [r"\bdefect\b", r"\bsnag\b", r"\brepair\b", r"\bAaron\b", r"\bpromise\b", r"\bassurance\b", r"\binspection\b", r"\bhandover\b"],
    }
    summary = {
        "files": file_info,
        "digest": str(OUT),
        "hits": {name: context_hits(all_text, pats) for name, pats in categories.items()},
    }
    SUMMARY.write_text(json.dumps(summary, indent=2), encoding="utf-8")
    print(json.dumps(file_info, indent=2))
    print(f"Digest: {OUT}")


if __name__ == "__main__":
    main()
