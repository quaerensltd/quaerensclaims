from pathlib import Path
from docx import Document

BASE = Path(r"C:\Users\CasaT\quaerensclaims\document_work\caravan_claim_review")
OUT = BASE / "caravan_bundle_text_digest.txt"

def docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [" ".join(c.text.split()) for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts)

with OUT.open("w", encoding="utf-8") as f:
    for path in sorted(BASE.glob("*.docx")):
        f.write("\n" + "=" * 90 + "\n")
        f.write(path.name + "\n")
        f.write("=" * 90 + "\n")
        try:
            f.write(docx_text(path))
        except Exception as exc:
            f.write(f"[EXTRACTION FAILED: {exc}]")
        f.write("\n")

print(OUT)
