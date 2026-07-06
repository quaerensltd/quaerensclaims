from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\CasaT\quaerensclaims\document_work\caravan_claim_review\Kennedy_Vaughan_Haven_Caravan_Assessment_Report.docx")

BLUE = "1F4E79"
LIGHT_BLUE = "EAF2FB"
LIGHT_GREY = "F2F4F7"
GREEN = "0B7A3B"
RED = "9B1C1C"
GOLD = "7A5A00"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.color.rgb = RGBColor.from_string(BLUE)
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = RGBColor.from_string(BLUE)
            run.font.size = Pt(13)
        else:
            run.font.color.rgb = RGBColor.from_string("1F4D78")
            run.font.size = Pt(12)
    return p


def add_para(doc, text, bold_start=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.add_run(text)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="000000")
        shade_cell(hdr[i], LIGHT_GREY)
        if widths:
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.size = Pt(11)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(3)
    for run in p2.runs:
        run.font.size = Pt(10)
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.1

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Initial Caravan Claim Assessment Report")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor.from_string(BLUE)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Kennedy Vaughan v Haven Holidays Limited / Bourne Leisure Ltd")
r.bold = True
r.font.size = Pt(12)

meta_rows = [
    ("Prepared for", "Quaerens initial review"),
    ("Prepared date", "3 July 2026"),
    ("Claim type", "Static holiday caravan purchase / alleged misrepresentation, defects and unresolved complaint"),
    ("Documents reviewed", "Haven litigation bundle draft, sections 1-20, master chronology and related draft bundle files"),
    ("Evidence status", "Narrative bundle reviewed; original exhibits, invoices, photographs and full dated correspondence still need final indexing"),
]
add_table(doc, ["Item", "Details"], meta_rows, [Inches(1.65), Inches(4.75)])

add_callout(
    doc,
    "Executive assessment",
    "On the documents reviewed, the case appears to have a potentially arguable basis against Haven, principally around alleged pre-contract assurances that identified caravan defects would be rectified before handover, reliance on those assurances, continuing defects after handover, and a prolonged unresolved complaint history. The case is not yet ready to quantify robustly because the bundle repeatedly states that exact dates, exhibit references and loss figures remain to be inserted.",
)

add_heading(doc, "1. Case Snapshot", 1)
snapshot_rows = [
    ("Claimant", "Kennedy Vaughan"),
    ("Potential respondent", "Haven Holidays Limited / Bourne Leisure Ltd - exact contractual entity to be confirmed from sale documents"),
    ("Product / service", "Static holiday caravan purchase and associated ownership services"),
    ("Core allegation", "The claimant says Haven induced the replacement caravan purchase by promising that pre-identified defects would be rectified before handover, but those works were not completed."),
    ("Claim status", "A litigation bundle has been drafted, but final exhibits, exact dates and a quantified Schedule of Loss remain incomplete."),
]
add_table(doc, ["Field", "Current position"], snapshot_rows, [Inches(1.8), Inches(4.6)])

add_heading(doc, "2. Documents Reviewed", 1)
add_para(doc, "The review is based on the extracted content from the following document bundle:")
for item in [
    "Draft Litigation Bundle (Progress Version)",
    "Master Bundle V2.1 / Master Chronology",
    "Sections 1-6 concerning background, parties, family circumstances, first caravan and upgrade",
    "Sections 8-20 concerning representations, defects, Consumer Rights Act issues, loss, SAR disclosure and misrepresentation",
    "Section 7c nested web/export files were inspected at folder level; these appear to contain web/chat asset files rather than primary claim evidence.",
]:
    add_bullet(doc, item)

add_heading(doc, "3. Factual Summary", 1)
for para in [
    "Mr Vaughan and his wife purchased a static holiday caravan from Haven for family use, not as an investment or commercial letting business. The bundle explains that the purchase was made in difficult family circumstances, including their son's lymphoma treatment and Mr Vaughan's own history of brain aneurysms and memory issues.",
    "After the first caravan purchase, Haven representatives allegedly encouraged an upgrade to a replacement caravan. The key Haven representative is identified as Aaron. The bundle says Aaron knew the family's circumstances and that a relationship of trust had developed over time.",
    "Before signing for the replacement caravan, Mr and Mrs Vaughan allegedly inspected it and identified numerous defects. The bundle says Aaron repeatedly assured them that every identified defect would be rectified before handover.",
    "The claimant says the purchase proceeded because those assurances were accepted and relied upon. After handover, the promised condition was allegedly not achieved. The bundle refers to approximately seventy-one defects, repeated complaints, repair requests, complaint escalation, a Subject Access Request and eventual preparation for litigation.",
]:
    add_para(doc, para)

add_heading(doc, "4. Potential Legal and Complaint Issues", 1)
issue_rows = [
    ("Misrepresentation / inducement", "Potentially strong if the claimant can evidence clear pre-contract statements, reliance and that the caravan was not delivered in the promised condition."),
    ("Breach of contract", "Potentially arguable if the promise to complete remedial works became part of the agreement or can be evidenced as a binding commitment."),
    ("Consumer Rights Act 2015", "Potentially arguable where goods supplied by a trader to a consumer were not of satisfactory quality, fit for purpose, as described, or failed to match pre-contract information. Specific remedies and limitation require legal review."),
    ("Complaint handling / delay", "Useful background evidence. It may support conduct, reasonableness and chronology, but is less likely to be the main cause of action unless specific legal duties or losses are evidenced."),
    ("Subject Access Request delay", "Relevant background and may support transparency concerns. Any data protection complaint should be treated separately from the core caravan sale claim."),
]
add_table(doc, ["Issue", "Assessment"], issue_rows, [Inches(2.05), Inches(4.35)])

add_heading(doc, "5. Strengths Identified", 1)
for item in [
    "The bundle has a consistent central narrative: defects were known before purchase, assurances were given, and the claimant relied on them.",
    "The claimant appears to have taken steps to inspect the caravan before purchase, which helps show that the concern was condition-specific rather than general dissatisfaction.",
    "The documents repeatedly refer to contemporaneous sources: emails, photographs, snagging records, complaint correspondence, internal Haven records and SAR disclosure.",
    "The claimant appears to have given Haven repeated opportunities to remedy the defects before litigation was pursued.",
    "The alleged personal circumstances may help explain reliance, vulnerability, trust and the practical impact of the dispute.",
]:
    add_bullet(doc, item)

add_heading(doc, "6. Weaknesses and Evidence Gaps", 1)
for item in [
    "The bundle is still a draft and repeatedly states that exact dates, exhibit references and page references remain to be inserted.",
    "Purchase price, finance route, site fees, insurance, removal costs and other loss figures are not presently quantified in the extracted bundle.",
    "The actual written purchase agreement, terms and conditions, warranty terms and any disclaimers must be reviewed before liability can be assessed properly.",
    "The alleged oral assurances must be supported by contemporaneous notes, emails, messages, snagging lists, witness statements or Haven records where possible.",
    "Limitation needs urgent legal review because some events appear to date back to 2016/2017 or later, and the exact upgrade date is not clear from the extracted bundle.",
    "If the caravan has now been removed, evidence is needed showing when, why, who authorised it, residual value, sale/removal proceeds if any, and how the removal affected losses.",
]:
    add_bullet(doc, item)

add_heading(doc, "7. Preliminary Merits View", 1)
add_table(
    doc,
    ["Area", "Rating", "Reason"],
    [
        ("Liability narrative", "Moderate to strong", "The alleged facts form a coherent claim if supported by exhibits."),
        ("Evidence readiness", "Moderate", "The bundle references evidence but does not yet fully index or attach it in the extracted material."),
        ("Quantum readiness", "Weak to moderate", "Loss categories are described, but figures are mostly TBC."),
        ("Limitation / procedural risk", "High risk until reviewed", "Dates are not finalised and the dispute appears historic."),
        ("Suitability for legal partner review", "Yes", "The matter is fact-heavy, potentially litigated and should be reviewed by a solicitor once the evidence index is complete."),
    ],
    [Inches(1.8), Inches(1.35), Inches(3.25)],
)

add_heading(doc, "8. Losses to Build Into a Schedule", 1)
add_para(doc, "The bundle suggests the following loss heads should be checked and evidenced. These should not be advanced unless documentary support is available.")
for item in [
    "Purchase price of the replacement caravan and any trade-in or part-exchange allowance.",
    "Annual site fees, pitch fees and insurance paid during the affected period.",
    "Costs of repair, inspection, removal, storage, transport or disposal.",
    "Any finance interest, charges or settlement costs if the caravan was financed.",
    "Loss of use, loss of enjoyment and inconvenience, subject to legal advice on recoverability.",
    "Complaint, correspondence and document request costs where properly recoverable.",
]:
    add_bullet(doc, item)

add_heading(doc, "9. Recommended Next Steps", 1)
steps = [
    "Confirm the correct legal entity from the purchase agreement and any Haven/Bourne Leisure paperwork.",
    "Create a dated chronology with exact dates for first purchase, upgrade, inspection, handover, first complaint, repair visits, formal complaint, SAR, disclosure, removal and solicitor involvement.",
    "Attach and label every key exhibit: sale agreement, terms, invoice, finance agreement, snagging list, photographs, emails, WhatsApp messages, complaint letters, Haven responses and SAR disclosure.",
    "Prepare a quantified Schedule of Loss with supporting proof for each figure.",
    "Obtain a signed witness statement from Mr Vaughan and, if possible, Mrs Vaughan covering the inspection, assurances, reliance and impact.",
    "Have limitation, contract terms and prospects reviewed by a litigation solicitor before threatening proceedings.",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    p.add_run(step)

add_heading(doc, "10. Overall Conclusion", 1)
add_para(
    doc,
    "Initial view: ",
    bold_start="Initial view: ",
)
doc.paragraphs[-1].add_run(
    "this appears to be a potentially worthwhile caravan claim for further legal review, provided the promised-remedial-works evidence and loss figures can be tied to documents. The strongest route is not simply that the caravan had defects; it is that defects were allegedly identified before purchase, Haven allegedly promised rectification before handover, the claimant relied on that promise, and the delivered position did not match what was represented."
)
add_para(
    doc,
    "The case should not be presented as fully ready until the exact contractual entity, dates, exhibits, limitation position and Schedule of Loss are completed."
)

add_callout(
    doc,
    "Important limitation",
    "This report is an initial evidence assessment based on the documents provided. It is not legal advice and does not guarantee recovery. Given the historic nature of the events and the fact that litigation is contemplated, the matter should be reviewed by an appropriately qualified legal partner before formal proceedings or settlement demands are issued.",
    fill="FFF2CC",
)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Quaerens initial case assessment - confidential working document")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor.from_string("666666")

doc.save(OUT)
print(OUT)
