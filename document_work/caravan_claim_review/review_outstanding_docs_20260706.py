from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Users\CasaT\quaerensclaims\document_work\caravan_claim_review\outstanding_docs_20260706")
OUT_TXT = ROOT.parent / "outstanding_docs_20260706_digest.txt"
OUT_JSON = ROOT.parent / "outstanding_docs_20260706_findings.json"


KEYWORDS = {
    "client_details": ["kennedy vaughan", "nicola vaughan", "1 forder grove", "b14 5jn", "telephone", "email"],
    "caravan_details": ["regal", "symphony", "cris", "pitch", "plot", "31", "shorfield", "hopton"],
    "first_purchase": ["purchase", "first caravan", "initial caravan", "sales agreement", "order form"],
    "upgrade_purchase": ["upgrade", "replacement caravan", "second caravan", "part exchange", "part-exchange"],
    "finance": ["evergreen", "finance agreement", "credit agreement", "settlement", "default", "termination"],
    "payment_proof": ["bank statement", "card payment", "credit card", "bacs", "receipt", "payment"],
    "site_fees": ["pitch fee", "site fee", "rent", "annual", "ledger"],
    "decking_extras": ["decking", "extras", "gas", "electricity", "water", "rates"],
    "snagging": ["snagging", "defect", "repair", "maintenance", "work order", "engineer"],
    "handover": ["handover", "pre-handover", "inspection", "occupation"],
    "promises": ["promise", "represented", "assured", "told", "representation"],
    "complaints": ["complaint", "response", "final response", "crm"],
    "sar": ["subject access", "sar", "disclosure", "late disclosure"],
    "vulnerability": ["aneurysm", "brain", "health", "son", "family circumstances", "vulnerable"],
    "family_deposits": ["family", "friends", "deposit", "returned", "holiday"],
    "loss_schedule": ["schedule of loss", "financial losses", "loss", "statutory interest"],
    "desired_outcome": ["desired outcome", "compensation", "statutory interest", "legal costs"],
    "limitation_dates": ["limitation", "chronology", "timeline", "2017", "2018", "2019"],
    "prior_action": ["court", "solicitor", "ombudsman", "settlement", "legal proceedings"],
    "authority": ["letter of authority", "authority", "authorise", "consent"],
}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def read_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        return f"[DOCX READ ERROR: {exc}]"


def find_matches(text: str) -> dict[str, list[str]]:
    low = text.lower()
    matches: dict[str, list[str]] = {}
    for category, words in KEYWORDS.items():
        hits = [w for w in words if w in low]
        if hits:
            matches[category] = hits
    return matches


def snippet_for(text: str, terms: list[str], width: int = 220) -> str:
    low = text.lower()
    best = None
    for term in terms:
        idx = low.find(term.lower())
        if idx >= 0 and (best is None or idx < best):
            best = idx
    if best is None:
        return ""
    start = max(0, best - width // 2)
    end = min(len(text), best + width)
    return re.sub(r"\s+", " ", text[start:end]).strip()


def main() -> None:
    files = [p for p in ROOT.rglob("*") if p.is_file()]
    records = []
    seen: dict[str, str] = {}

    for path in files:
        digest = sha256(path)
        duplicate_of = seen.get(digest)
        if duplicate_of is None:
            seen[digest] = str(path)
        text = ""
        if path.suffix.lower() == ".docx":
            text = read_docx(path)
        elif path.suffix.lower() == ".pdf":
            text = "[PDF file present - text not extracted in this pass]"
        elif path.suffix.lower() in {".png", ".jpg", ".jpeg"}:
            text = "[Image file present - visual evidence only]"

        matches = find_matches(text + "\n" + path.name)
        records.append(
            {
                "path": str(path),
                "name": path.name,
                "suffix": path.suffix.lower(),
                "size": path.stat().st_size,
                "sha256": digest,
                "duplicate_of": duplicate_of,
                "text_chars": len(text),
                "matches": matches,
                "snippets": {
                    cat: snippet_for(text, terms)
                    for cat, terms in matches.items()
                    if text and not text.startswith("[Image")
                },
            }
        )

    by_category: dict[str, list[dict[str, str]]] = {k: [] for k in KEYWORDS}
    for rec in records:
        if rec["duplicate_of"]:
            continue
        for cat in rec["matches"]:
            by_category[cat].append(
                {
                    "name": rec["name"],
                    "path": rec["path"],
                    "snippet": rec["snippets"].get(cat, ""),
                }
            )

    with OUT_JSON.open("w", encoding="utf-8") as f:
        json.dump({"records": records, "by_category": by_category}, f, indent=2)

    lines: list[str] = []
    lines.append(f"Files scanned: {len(records)}")
    lines.append(f"Unique files by SHA256: {len(seen)}")
    lines.append("")
    lines.append("EXTENSIONS")
    ext_counts: dict[str, int] = {}
    for rec in records:
        ext_counts[rec["suffix"] or "[no ext]"] = ext_counts.get(rec["suffix"] or "[no ext]", 0) + 1
    for ext, count in sorted(ext_counts.items()):
        lines.append(f"- {ext}: {count}")
    lines.append("")
    lines.append("CATEGORY MATCHES")
    for cat, hits in by_category.items():
        lines.append(f"\n## {cat} ({len(hits)} unique files)")
        for hit in hits[:12]:
            lines.append(f"- {hit['name']}")
            if hit["snippet"]:
                lines.append(f"  {hit['snippet'][:450]}")
    OUT_TXT.write_text("\n".join(lines), encoding="utf-8")
    print(OUT_TXT)
    print(OUT_JSON)


if __name__ == "__main__":
    main()
