from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\CasaT\quaerensclaims\document_work\caravan_claim_review\01_Kennedy_Vaughan_Haven_Caravan_Assessment_Report_Professional_Updated_v3.docx")

BLUE = "005BAA"
DARK = "0B1F3A"
LIGHT_BLUE = "EAF4FF"
PALE = "F7FBFF"
GREY = "F2F4F7"
GOLD = "FFF2CC"
GREEN = "0B7A3B"
RED = "9B1C1C"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def borders(cell, color="D9E5F2", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        elem = tc_borders.find(qn("w:" + edge))
        if elem is None:
            elem = OxmlElement("w:" + edge)
            tc_borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def cell_text(cell, text, bold=False, size=9.5, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    borders(cell)


def para(doc, text=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    p.add_run(text)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else DARK)
        r.font.size = Pt(15 if level == 1 else 12)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    p.add_run(text)


def simple_table(doc, headers, rows, widths=None, header_fill=GREY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell_text(cell, h, bold=True, size=9)
        shade(cell, header_fill)
        if widths:
            cell.width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell_text(cells[i], val, size=9)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def note_box(doc, text, title=None, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    borders(cell, color="BBD7F2", size="10")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    if title:
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(BLUE)
        p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    p.add_run(text).font.size = Pt(10)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(7)

    title_tbl = doc.add_table(rows=1, cols=1)
    title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = title_tbl.cell(0, 0)
    shade(cell, DARK)
    borders(cell, color=DARK, size="12")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("QUAERENS LTD\nUpdated Professional Client Assessment Report\nStatic Caravan / Holiday Park Claim Review")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("FFFFFF")
    r.font.size = Pt(16)
    doc.add_paragraph()

    top = doc.add_table(rows=1, cols=2)
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in top.rows[0].cells:
        shade(c, PALE)
        borders(c)
    cell_text(top.cell(0, 0), "Client\nKennedy Vaughan / Nicola Vaughan\n1 Forder Grove\nBirmingham\nWest Midlands\nB14 5JN", size=10)
    cell_text(top.cell(0, 1), "Prepared for\nUpdated claim assessment and evidence review\nDate prepared: 6 July 2026\nMatter: Haven / Hopton static caravan dispute", size=10)
    doc.add_paragraph()

    note_box(
        doc,
        "This is an updated assessment prepared from the client bundle, the further 6 July documentation response bundle, Hopton SAR workbook, insurance schedule, complaint material and schedule-of-loss drafts. It is not legal advice and does not guarantee recovery. The matter should be reviewed by an appropriate legal partner before any formal proceedings are issued.",
        "Important status note:",
        GOLD,
    )

    heading(doc, "1. Evidence Newly Reviewed")
    simple_table(
        doc,
        ["Source", "What it adds"],
        [
            ("16 new zip files", "193 Word files extracted; 98 unique document texts after duplicates. Mostly trial-bundle drafts, witness drafts, chronology, indexes, issue lists, relief and handover checklists."),
            ("SAR Hopton workbook", "25 sheets of Haven account, contact, complaint, finance, payment and activity records."),
            ("Insurance PDF", "Confirms insurance schedule information, caravan make/model and plot number."),
            ("Complaint to Hopton document", "Adds the client's own narrative, purchase/upgrade figures, snagging history, site fee dispute and desired resolution background."),
            ("Latest 13 client documents", "Adds focused statements on Haven's knowledge, reliance, promises, human impact, complaints, snagging and an evidence-based Schedule of Loss."),
            ("6 July outstanding documentation response bundle", "Five additional zip files reviewed. 196 files were scanned, with 124 unique files after duplicates. The bundle adds master witness statement sections, court chronology, limitation timeline, Letter of Authority, exhibit/document registers, SAR-based payment schedules, caravan identification records and litigation-pack indexes."),
            ("Hopton PDF", "Appears image-only/scanned on text extraction; no reliable text could be extracted without OCR or manual visual review."),
        ],
        [Inches(1.8), Inches(4.9)],
    )

    heading(doc, "2. Client and Account Information")
    simple_table(
        doc,
        ["Field", "Confirmed / evidenced information"],
        [
            ("Client names", "Kennedy Vaughan and Nicola Vaughan appear in Haven/SAR records. The insurance schedule names Mrs Nicola Vaughan with Mr Vaughan as additional policy holder."),
            ("Address", "1 Forder Grove, Birmingham, West Midlands, B14 5JN."),
            ("Telephone records", "Haven/SAR records show telephone entries ending 7772237667 and later 7547917159. Full client-confirmed current number remains to be checked."),
            ("Emails seen", "nicolaemmavaughan@icloud.com, cliffordcox@blueyonder.co.uk, kennedy.vaughan@aseeltd.com and ken@volgainvestmentltd.com appear in disclosed records."),
            ("Haven account", "Hopton Holiday Village account number HO055639."),
            ("Park", "Hopton Holiday Village, Great Yarmouth, Norfolk, NR31 9BW."),
            ("Pitch / plot", "Plot number 31 is confirmed by the insurance schedule. Complaint letter refers to pitch 31 Shorfield."),
            ("SAR stock records", "The 6 July caravan identification register records original stock number XC8881 and replacement stock number XA5988."),
            ("Purchase dates", "The 6 July purchase and caravan identification records state original purchase on 01 May 2017 and replacement/upgrade on 07 May 2018."),
            ("Caravan details", "Insurance schedule lists make Regal and model Symphony. The 6 July caravan identification register states that make, model year, CRiS number, serial number and manufacturer were not disclosed in the current SAR."),
            ("Insurance", "Policy number CETASLH/1267420. Underwritten by AIG Europe Limited via Towergate. Renewal date 31/05/2019. Total premium GBP 210.39."),
            ("Authority", "A Letter of Authority and Instructions to Quaerens has now been supplied, authorising disclosure from Haven, finance provider, insurer, warranty provider, contractors and other relevant parties. The telephone number/date fields should still be completed before use."),
        ],
        [Inches(1.65), Inches(5.05)],
    )

    heading(doc, "3. Executive Summary")
    para(doc, "The new material materially improves the factual basis of the report. It confirms the Hopton account, address, contact history, caravan make/model, plot number, complaint history, finance involvement and several payment/account entries. The latest client documents also strengthen the evidence around reliance, Haven's alleged knowledge of the family's circumstances, the snagging position, complaints and human impact.")
    para(doc, "The core complaint remains that Mr and Mrs Vaughan purchased and then upgraded a static caravan at Haven Hopton after relying on representations and assurances that identified defects and snagging would be resolved. Their own complaint letter says the second caravan was an older three-bedroom van, that they stayed on pitch 31 Shorfield, and that they were told to make a list of snags which would be sorted within three weeks.")
    para(doc, "The complaint letter states that more than 80 snags were listed, that many remained unresolved over two seasons, that the caravan was not fit for purpose, and that the family lost the practical enjoyment of the caravan during a period when their son was seriously ill. Haven's own SAR workbook records a complaint created on 30 July 2018 and later reopening/closure activity in November 2019.")
    para(doc, "The 6 July bundle gives the matter a more organised litigation-style structure. It supports the account, stock, purchase/upgrade, finance, payment and pitch-fee chronology, and it includes a Letter of Authority. It also confirms that several key original documents remain pending, draft or still to be populated, so the matter is stronger but not yet complete.")

    note_box(
        doc,
        "Preliminary strength: stronger than before. The case now has documentary support for identity, account, park, complaint dates, finance records, stock numbers, payment registers, snagging position and SAR-backed loss indicators. The main remaining risk is still legal review, limitation and proof: exact contract terms, final disposal/refund figures, finance settlement documents, original bank/receipt evidence and exhibit-level proof are still needed.",
        "Updated assessment:",
        LIGHT_BLUE,
    )

    heading(doc, "4. What the 6 July Bundle Now Adds")
    simple_table(
        doc,
        ["Area", "Updated evidence position"],
        [
            ("Evidence structure", "The response bundle contains 196 files, 124 unique by hash after duplicates. It includes master witness statement sections, chronology, limitation timeline, evidence/exhibit registers, source registers and volume-based document schedules."),
            ("Client authority", "A Letter of Authority and Instructions has been prepared for Quaerens/solicitors, covering Haven, finance provider, insurer, warranty provider, contractors, engineers, surveyors and other relevant third parties. It should be completed with missing telephone/date details before being sent."),
            ("Account and caravan records", "Account HO055639, original stock XC8881 and replacement stock XA5988 are recorded. The bundle records original purchase on 01 May 2017 and replacement on 07 May 2018."),
            ("Finance", "The bundle records Evergreen Finance Ltd and finance agreement 12022646, with finance advance GBP 33,750, replacement transaction GBP 29,000 and finance settlement GBP 26,220.83. Earlier material also referred to agreement 12022154, so agreement numbers should be reconciled against the original finance documents."),
            ("Payment schedule", "Payment registers record GBP 3,000, GBP 2,000, GBP 5,000, GBP 7,500 and GBP 4,000 payments, plus SAR-backed finance and ledger entries. Original bank/card statements and merchant receipts are still needed to make the loss schedule court-ready."),
            ("Pitch and site charges", "The bundle records 2017 annual site fees of GBP 5,863 and 2018 annual site fees of GBP 6,150, supported by SAR ledger-style records. Utility, rates and insurance entries are also identified."),
            ("Snagging and defects", "The witness/snagging material refers to approximately 71 defects and a substantial number still outstanding. The master snagging schedule remains draft and needs final exhibit references, photographs and repair/inspection evidence."),
            ("Schedule of loss", "A reconciled schedule exists, but it remains a working draft. It specifically says purchase contracts, upgrade documents, finance agreements, bank statements, decking/veranda invoice and final owner-leaving documents are still required."),
        ],
        [Inches(1.8), Inches(4.9)],
    )

    heading(doc, "5. Key Chronology")
    simple_table(
        doc,
        ["Date / period", "Event / evidence"],
        [
            ("28 Sep 2016", "Haven/SAR record created for Mrs Nicola Vaughan at 1 Forder Grove, Birmingham."),
            ("14 Apr 2017", "Sales/contact notes record Mr Vaughan looking to purchase, with family circumstances and finance discussion noted."),
            ("29 Apr 2017", "Hopton appointment recorded as made and kept."),
            ("1 May 2017", "Haven records show sale activity and card payments of GBP 3,000 and GBP 2,000."),
            ("20 May 2017", "Finance/deposit entries for sale HO00B8BE: Evergreen Finance Ltd, agreement 12022154, finance GBP 33,750, deposit GBP 5,000 and due later GBP 7,500."),
            ("30/31 May 2017", "Records show BACS payment GBP 7,500 and finance received GBP 33,750."),
            ("2017 ownership period", "Client complaint says first purchase was approximately GBP 48,000 and was initially satisfactory but too small for family needs."),
            ("May 2018", "Upgrade/sale records show finance entries around GBP 29,000 and a GBP 2,000 card payment. Complaint letter states the second upgrade was approximately GBP 61,000."),
            ("June 2018", "Haven records show finance received GBP 29,000 for Evergreen agreement 12022646."),
            ("30 Jul 2018", "Haven/SAR records show complaint created/received and emails acknowledged. Complaint subject: owner part-exchange June 2018, many snagging issues not resolved, alleged breach and not fit for purpose."),
            ("2018 season", "Client says more than 80 snags were listed and many remained unresolved, causing the family to stay away while full site fees and bills were paid."),
            ("7 Feb 2019", "Haven/SAR records show GBP 4,000 credit card payment received. Client says this was paid while withholding the balance until snags were completed."),
            ("1 Aug 2019", "Haven/SAR records contain termination letter for overdue balance GBP 4,657.15 and letter to Evergreen Finance confirming termination and storage/removal position."),
            ("8 Sep 2019", "Complaint letter says defects remained, disconnection charge had been applied, services remained connected and the client attended to clear site."),
            ("12 Nov 2019", "Haven/SAR records show complaint closed. Internal note says complaint closed/no response sent because Gareth B was dealing with it by email."),
            ("19 Nov 2019", "Haven/SAR financial records show OLP entries including finance company payment GBP 26,220.83, caravan debtor rent GBP -44,652.63 and rent cheque payment GBP 13,671.11."),
            ("26 Nov 2019", "Haven/SAR email explains rent ledger balance of GBP 4,503.91 and refers to refund shown on owner leaving park form."),
            ("2025/2026", "Later SAR bundle indicates subject access and disclosure activity, with delayed disclosure concerns already captured in the prior report."),
        ],
        [Inches(1.55), Inches(5.15)],
    )

    heading(doc, "6. Financial and Loss Indicators")
    para(doc, "The new files now include both a general Schedule of Loss and an evidence-based Schedule of Loss prepared from SAR data. It remains a working schedule, but it gives a much clearer starting point for legal review.")
    simple_table(
        doc,
        ["Loss / figure", "Evidence / status"],
        [
            ("First caravan purchase", "Client complaint states approximately GBP 48,000. Needs contract/invoice confirmation."),
            ("Part exchange / value returned", "Complaint says Julia Mills gave GBP 29,000 back, leaving the client feeling tied into an upgrade and losing nearly GBP 20,000. Needs supporting sales/part-exchange documents."),
            ("Second upgrade", "Complaint states approximately GBP 61,000 for the older three-bedroom replacement caravan."),
            ("SAR schedule entries", "The latest evidence-based Schedule of Loss records GBP 3,000 and GBP 2,000 card payments on 01/05/2017, GBP 5,000 deposit, GBP 33,750 finance and GBP 7,500 balance due later."),
            ("Further SAR-backed charges", "Schedule records GBP 6,331.62 card payment on 30/01/2018, annual site fees GBP 5,863, caravan insurance GBP 310.95, contents insurance GBP 33.88, electricity GBP 60.74, metered gas GBP 72.99, non-domestic rates GBP 220.91 and water/sewerage GBP 437.10."),
            ("Site fees / bills", "Complaint states over GBP 7,000 paid for site fees including bills while the caravan was not fit for purpose. SAR ledger entries show annual site fees GBP 5,863 in 2017 and GBP 6,150 in 2018 plus utilities/rates."),
            ("2019 payment withheld/paid", "Client says GBP 4,000 was paid in January 2019; SAR records show GBP 4,000 card payment received on 7 February 2019."),
            ("Overdue balance / termination", "Termination letter in SAR record shows overdue balance GBP 4,657.15 on 1 August 2019."),
            ("Disconnection fee", "Complaint states Haven charged GBP 160; SAR ledger shows service/disconnection entry GBP 139.34 on 31 July 2019."),
            ("Decking/veranda", "Complaint says decking cost was over GBP 10,000 and that the client believed it belonged to them. Needs invoice/contract proof."),
            ("Finance settlement", "SAR/6 July finance records show Evergreen Finance Ltd agreement 12022646 and a later payment to finance company of GBP 26,220.83. Earlier records also refer to agreement 12022154, so the finance documents must reconcile the agreement history."),
            ("Refund / leaving park", "SAR email refers to refund on owner leaving park form and rent ledger balance GBP 4,503.91. Owner leaving park form and cheque/refund detail still needed."),
        ],
        [Inches(1.75), Inches(4.95)],
    )

    heading(doc, "7. Core Liability Themes")
    for item in [
        "Misrepresentation / reliance: the client says he was told snagging would be completed within three weeks and relied on that assurance when proceeding.",
        "Consumer quality / fitness: the complaint says the replacement caravan was not fit for purpose because more than 80 snags were listed and many remained unresolved.",
        "Failure to repair within a reasonable time: client narrative and SAR complaint history suggest repeated opportunities were given.",
        "Complaint handling: Haven records show complaint creation, acknowledgement, reopening and closure, including an internal note that no customer relations response was sent because another manager was dealing with it.",
        "Finance-linked consequences: Evergreen Finance Ltd and agreement numbers are now confirmed in the SAR records, which may be relevant to liability route and recoverability.",
        "Knowledge and vulnerability context: the latest documents state that Haven representatives were aware of Mr Vaughan's recovery from brain aneurysms and the family's reasons for buying the caravan, including their son's serious illness.",
    ]:
        bullet(doc, item)

    heading(doc, "8. Promises, Reliance and Snagging Evidence")
    para(doc, "The latest documents provide a more focused statement of what the claimant says was promised. The key alleged representations are that the replacement caravan would be prepared to an acceptable standard before handover, snagging works would be completed promptly, the caravan would be fit for occupation, any further defects would be addressed through Haven's maintenance process, and the upgrade would provide a better standard of accommodation and ownership experience.")
    para(doc, "The reliance position is now clearer. The claimant states that he proceeded because he trusted Haven's expertise and assurances, particularly given his personal and family circumstances at the time. Those facts may be relevant when considering whether the representations were material and whether the claimant's reliance was reasonable.")
    simple_table(
        doc,
        ["Snagging point", "Current evidence position"],
        [
            ("Approximate number of snagging items", "Latest snagging documents state approximately 71 snagging defects/items were identified."),
            ("Timing", "The documents state defects were identified during the pre-handover inspection before signing/taking possession."),
            ("Assurance alleged", "Haven representative allegedly assured the claimant the identified items would be completed before occupation or within a reasonable period."),
            ("Outstanding position", "The claimant states approximately 48 defects remained outstanding when the caravan was ultimately removed, and/or more than 60% remained unresolved when the family vacated/removed belongings."),
            ("Evidence still needed", "The detailed itemised snagging list, photographs, repair records, engineer reports and exhibit references still need to be indexed for solicitor review."),
        ],
        [Inches(2.15), Inches(4.55)],
    )

    heading(doc, "9. Human Impact and Loss of Enjoyment")
    para(doc, "The new human impact documents explain that the caravan was intended to provide a peaceful family place after Mr Vaughan's brain aneurysms and his son's serious illness. The client says Haven representatives were aware of those circumstances because they were openly discussed during the sales process.")
    para(doc, "The loss of use was not simply financial. The client says family and friends had planned holidays and, in some cases, paid deposits, which he then had to return. The documents describe embarrassment, stress, damaged relationships, loss of confidence and the emotional burden of feeling that he had failed people close to him.")
    para(doc, "These matters should be used carefully. They strengthen the context and loss-of-enjoyment narrative, but any formal claim for distress, inconvenience or consequential loss should be reviewed by a legal partner and supported only where recoverable.")

    heading(doc, "10. Desired Outcome")
    para(doc, "Mr Vaughan's primary objective is to obtain a fair and lawful resolution to the dispute with Haven Holidays Ltd. He seeks appropriate compensation for the financial losses suffered as a result of the purchase and subsequent upgrade of the holiday caravan, together with any additional losses that are legally recoverable.")
    para(doc, "He also seeks recognition of the significant inconvenience, distress and prolonged period during which he attempted to resolve the matter directly with Haven before considering legal proceedings. His preference has been to resolve the dispute without litigation, but despite repeated complaints and opportunities for the matter to be addressed, this has not been achieved.")
    for item in [
        "Recovery of financial losses established by the evidence.",
        "Statutory interest where recoverable.",
        "Recovery of legal costs where a court or applicable process has jurisdiction to award them.",
        "Any other remedy considered fair and appropriate based on the evidence.",
        "An independent and impartial assessment of the facts so that the dispute can be resolved fairly and finally.",
    ]:
        bullet(doc, item)

    heading(doc, "11. Information Still Needed")
    simple_table(
        doc,
        ["Outstanding item", "Why it matters"],
        [
            ("Full contract pack for both purchases/upgrades", "Needed to prove exact purchase terms, price, warranties, part-exchange allowance and any exclusions."),
            ("Finance agreements and settlement statements", "Evergreen agreement 12022646 is identified in the 6 July bundle and earlier material also refers to 12022154. The full agreements, payment history and final settlement statements are required."),
            ("Itemised snagging list and photographs", "The new documents state approximately 71 items and 48 outstanding, but the exact itemised list and image evidence should be exhibited with references."),
            ("Written record of assurances", "The representation statement is helpful, but any email/text/note from Aaron or Haven promising completion within three weeks or before handover remains high priority."),
            ("Owner leaving park form / final account", "SAR and the 6 July loss schedule mention refund/leaving park figures, including GBP 4,503.91. The original owner leaving park form and payment/refund confirmation are needed to calculate net losses."),
            ("Decking/veranda invoice and ownership terms", "The complaint states over GBP 10,000 was spent and ownership/removal was disputed."),
            ("Evidence of full refund offer", "Complaint says a full refund was offered at the start of the season. The exact offer should be located."),
            ("Current client contact confirmation", "SAR contains several phone/email records. Current preferred address, phone and email should be confirmed before correspondence."),
            ("Medical/family impact evidence if relied upon", "The new human impact statements are helpful, but only relevant, proportionate supporting evidence should be used if distress/loss of enjoyment is claimed."),
            ("Limitation review", "The key events are historic. A legal partner must assess limitation, concealment/knowledge arguments and the correct cause of action."),
        ],
        [Inches(2.2), Inches(4.5)],
    )

    heading(doc, "12. Recommended Next Steps")
    for item in [
        "Ask the client to confirm whether the current primary claimant should be Mr Kennedy Vaughan, Mrs Nicola Vaughan, or both.",
        "Obtain the first and second purchase agreements, finance agreements, sales invoices, part-exchange documents and owner leaving park form.",
        "Create a final Schedule of Loss using the SAR ledger, complaint figures, finance settlement, refund/cheque details, decking cost and site fees.",
        "Prepare a clean exhibit index: contract documents, SAR workbook extracts, complaint letter, insurance schedule, snagging list, photos, ledger and termination letters.",
        "Refer the completed evidence pack to a legal partner for limitation, liability and recoverability advice before sending any formal letter of claim.",
    ]:
        bullet(doc, item)

    note_box(
        doc,
        "This updated report is stronger than the previous version because several details previously marked TBC are now supported by the Hopton SAR workbook, insurance schedule and the latest focused client statements. The case should still not be treated as complete: contract documents, finance documents, final account/refund evidence, and exhibit-level proof remain essential.",
        "Conclusion:",
        LIGHT_BLUE,
    )

    para(doc, "Disclaimer: This report is an initial assessment only. It is based on documents provided and extracted as at the report date. It does not constitute legal advice, does not confirm that any claim will succeed and should be reviewed by a suitably authorised legal professional before formal action.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
