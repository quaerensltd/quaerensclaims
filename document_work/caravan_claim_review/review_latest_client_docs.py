from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document


WORK = Path(r"C:\Users\CasaT\quaerensclaims\document_work\caravan_claim_review")
DIGEST = WORK / "latest_client_docs_digest.txt"
SUMMARY = WORK / "latest_client_docs_findings.json"

FILES = [
    Path(r"C:\Users\CasaT\Downloads\loss of waggon earnings .docx"),
    Path(r"C:\Users\CasaT\Downloads\Haven knowledge of circumstances .docx"),
    Path(r"C:\Users\CasaT\Downloads\Quaerens_Assessment_Information_Sheet.docx"),
    Path(r"C:\Users\CasaT\Downloads\assessment .docx"),
    Path(r"C:\Users\CasaT\Downloads\reliance on haven .docx"),
    Path(r"C:\Users\CasaT\Downloads\human impact .docx"),
    Path(r"C:\Users\CasaT\Downloads\impact & dispute .docx"),
    Path(r"C:\Users\CasaT\Downloads\what was promised by haven .docx"),
    Path(r"C:\Users\CasaT\Downloads\snagging list +.docx"),
    Path(r"C:\Users\CasaT\Downloads\snagging list .docx"),
    Path(r"C:\Users\CasaT\Downloads\complaints .docx"),
    Path(r"C:\Users\CasaT\Downloads\00.7_Schedule_of_Loss_V1+.docx"),
    Path(r"C:\Users\CasaT\Downloads\00.7_Schedule_of_Loss_V1.docx"),
]


def read_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
    except Exception as exc:
        return f"[UNREADABLE DOCX: {exc}]"
    parts: list[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def context_hits(text: str, patterns: list[str], radius: int = 260) -> list[str]:
    hits: list[str] = []
    for pat in patterns:
        for match in re.finditer(pat, text, flags=re.IGNORECASE):
            start = max(0, match.start() - radius)
            end = min(len(text), match.end() + radius)
            snippet = re.sub(r"\s+", " ", text[start:end]).strip()
            if snippet not in hits:
                hits.append(snippet)
            if len(hits) >= 12:
                return hits
    return hits


def main() -> None:
    chunks: list[str] = []
    file_info: list[dict] = []
    for file in FILES:
        info = {"path": str(file), "exists": file.exists()}
        if file.exists():
            text = read_docx(file)
            info["chars"] = len(text)
            chunks.append(f"\n\n===== {file} =====\n{text}")
        file_info.append(info)

    all_text = "\n".join(chunks)
    DIGEST.write_text(all_text, encoding="utf-8", errors="ignore")

    categories = {
        "contact_and_assessment_sheet": [r"\baddress\b", r"\btelephone\b", r"\bemail\b", r"\bQuaerens\b", r"\bassessment\b"],
        "promises_reliance": [r"\bpromise", r"\bassurance", r"\breliance", r"\brelied", r"\bAaron\b", r"\bHaven knew\b"],
        "snagging_defects": [r"\bsnag", r"\bdefect", r"\bleak", r"\bboiler", r"\brepair", r"\bfit for purpose", r"\b80\b"],
        "losses": [r"£\s?\d", r"\bloss", r"\bearnings", r"\bsite fees", r"\bdeck", r"\brefund", r"\bcompensation", r"\bschedule of loss\b"],
        "complaints": [r"\bcomplaint", r"\bJuly 2018\b", r"\bNovember 2019\b", r"\bterminated", r"\bemail", r"\bresponse\b"],
        "human_impact": [r"\bdistress", r"\bstress", r"\bson", r"\bcancer", r"\bfamily", r"\benjoyment", r"\bimpact\b"],
    }
    findings = {
        "files": file_info,
        "digest": str(DIGEST),
        "hits": {name: context_hits(all_text, pats) for name, pats in categories.items()},
    }
    SUMMARY.write_text(json.dumps(findings, indent=2), encoding="utf-8")
    print(json.dumps(file_info, indent=2))
    print(f"Digest: {DIGEST}")
    print(f"Summary: {SUMMARY}")


if __name__ == "__main__":
    main()
