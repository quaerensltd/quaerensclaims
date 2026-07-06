from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from collections import defaultdict
from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Users\CasaT\quaerensclaims")
WORK = ROOT / "document_work" / "caravan_claim_review"
EXTRACT_DIR = WORK / "new_client_docs"
DIGEST = WORK / "new_client_docs_digest.txt"
SUMMARY = WORK / "new_client_docs_findings.json"

ZIPS = [
    Path(r"C:\Users\CasaT\Downloads\Re_ trile bunble  (1).zip"),
    Path(r"C:\Users\CasaT\Downloads\[No Subject].zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_ .zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (1).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (2).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (3).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (4).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (5).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (6).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (7).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (8).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (9).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (10).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (11).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (12).zip"),
    Path(r"C:\Users\CasaT\Downloads\Re_  (13).zip"),
]


def safe_name(path: Path) -> str:
    base = path.stem.strip() or "blank"
    base = re.sub(r"[^A-Za-z0-9._()-]+", "_", base).strip("._ ")
    return base or "blank"


def safe_member_filename(name: str, fallback: str) -> str:
    raw = Path(name.replace("\\", "/")).name or fallback
    cleaned = re.sub(r"[^A-Za-z0-9 ._()-]+", "_", raw).strip(" .")
    return cleaned or fallback


def extract_all() -> None:
    if EXTRACT_DIR.exists():
        shutil.rmtree(EXTRACT_DIR)
    EXTRACT_DIR.mkdir(parents=True, exist_ok=True)
    for index, zip_path in enumerate(ZIPS, start=1):
        if not zip_path.exists():
            continue
        target = EXTRACT_DIR / f"{index:02d}_{safe_name(zip_path)}"
        target.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(zip_path) as zf:
            for member_index, member in enumerate(zf.infolist(), start=1):
                if member.is_dir():
                    continue
                filename = safe_member_filename(member.filename, f"file_{member_index}")
                out = target / f"{member_index:03d}_{filename}"
                out.parent.mkdir(parents=True, exist_ok=True)
                with zf.open(member) as src, out.open("wb") as dst:
                    dst.write(src.read())

    # Extract nested ZIPs into sibling folders.
    for nested in list(EXTRACT_DIR.rglob("*.zip")):
        nested_target = nested.parent / f"nested_{safe_name(nested)}"
        nested_target.mkdir(parents=True, exist_ok=True)
        try:
            with zipfile.ZipFile(nested) as zf:
                for member_index, member in enumerate(zf.infolist(), start=1):
                    if member.is_dir():
                        continue
                    filename = safe_member_filename(member.filename, f"nested_{member_index}")
                    out = nested_target / f"{member_index:03d}_{filename}"
                    out.parent.mkdir(parents=True, exist_ok=True)
                    with zf.open(member) as src, out.open("wb") as dst:
                        dst.write(src.read())
        except zipfile.BadZipFile:
            pass


def docx_text(path: Path) -> str:
    try:
        doc = Document(str(path))
    except Exception as exc:
        return f"[UNREADABLE DOCX: {exc}]"
    bits: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            bits.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                bits.append(row_text)
    return "\n".join(bits)


def fingerprint(text: str) -> str:
    normalised = re.sub(r"\s+", " ", text).strip().lower()
    return hashlib.sha256(normalised.encode("utf-8", errors="ignore")).hexdigest()


def context_hits(text: str, patterns: list[str], radius: int = 220) -> list[str]:
    hits: list[str] = []
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            start = max(0, match.start() - radius)
            end = min(len(text), match.end() + radius)
            snippet = re.sub(r"\s+", " ", text[start:end]).strip()
            if snippet not in hits:
                hits.append(snippet)
            if len(hits) >= 10:
                return hits
    return hits


def main() -> None:
    extract_all()
    docx_files = sorted(EXTRACT_DIR.rglob("*.docx"))
    by_hash: dict[str, dict] = {}
    duplicates: defaultdict[str, list[str]] = defaultdict(list)
    digest_chunks: list[str] = []

    for path in docx_files:
        text = docx_text(path)
        fp = fingerprint(text)
        duplicates[fp].append(str(path))
        if fp not in by_hash:
            by_hash[fp] = {"path": str(path), "text": text}
            digest_chunks.append(f"\n\n===== {path} =====\n{text}")

    all_text = "\n".join(item["text"] for item in by_hash.values())
    DIGEST.write_text("\n".join(digest_chunks), encoding="utf-8", errors="ignore")

    categories = {
        "contact_or_identity": [
            r"\bKennedy\b", r"\bVaughan\b", r"\bemail\b", r"\bphone\b", r"\bmobile\b",
            r"\baddress\b", r"\bpostcode\b",
        ],
        "park_pitch_caravan": [
            r"\bHopton\b", r"\bHaven\b", r"\bpitch\b", r"\bplot\b", r"\bcaravan\b",
            r"\bABI\b", r"\bSwift\b", r"\bWillerby\b", r"\bmodel\b", r"\bserial\b", r"\bCRiS\b",
        ],
        "prices_losses": [
            r"£\s?\d", r"\bprice\b", r"\bdeposit\b", r"\bsite fee", r"\brefund\b",
            r"\bloss\b", r"\bdamages\b", r"\bcompensation\b", r"\binterest\b",
        ],
        "dates": [
            r"\b20\d{2}\b", r"\b2016\b", r"\b2017\b", r"\b2024\b", r"\b2025\b", r"\b2026\b",
            r"\bJanuary\b", r"\bFebruary\b", r"\bMarch\b", r"\bApril\b", r"\bMay\b", r"\bJune\b",
            r"\bJuly\b", r"\bAugust\b", r"\bSeptember\b", r"\bOctober\b", r"\bNovember\b", r"\bDecember\b",
        ],
        "promises_defects": [
            r"\bAaron\b", r"\bpromise", r"\bassurance", r"\bdefect", r"\bsnag", r"\brepair",
            r"\binspection\b", r"\bhandover\b", r"\boccup", r"\bdamp\b", r"\bleak\b", r"\bunsafe\b",
        ],
        "complaints_sar": [
            r"\bcomplaint\b", r"\bSAR\b", r"\bsubject access\b", r"\bDSAR\b", r"\bICO\b",
            r"\bdisclosure\b", r"\badmin error\b", r"\bGDPR\b",
        ],
        "relief_route": [
            r"\brelief\b", r"\bremedy\b", r"\brescission\b", r"\bMisrepresentation\b",
            r"\bConsumer Rights Act\b", r"\bCPR\b", r"\bCounty Courts Act\b", r"\bpre-action\b",
        ],
    }

    findings = {
        "zip_count": sum(1 for z in ZIPS if z.exists()),
        "docx_file_count": len(docx_files),
        "unique_docx_text_count": len(by_hash),
        "duplicate_groups": sum(1 for paths in duplicates.values() if len(paths) > 1),
        "digest": str(DIGEST),
        "hits": {name: context_hits(all_text, pats) for name, pats in categories.items()},
        "unique_files": [item["path"] for item in by_hash.values()],
    }
    SUMMARY.write_text(json.dumps(findings, indent=2), encoding="utf-8")
    print(json.dumps({k: findings[k] for k in ["zip_count", "docx_file_count", "unique_docx_text_count", "duplicate_groups", "digest"]}, indent=2))


if __name__ == "__main__":
    main()
