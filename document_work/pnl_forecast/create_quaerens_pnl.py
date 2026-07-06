from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "Quaerens_12_Month_P_and_L_Forecast_and_Collaboration_Proposal.docx"


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0, 82, 164)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=(255, 255, 255))
        hdr[i]._tc.get_or_add_tcPr().append(parse_shading("0052A4"))
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    return table


def parse_shading(fill):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    return parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), fill))


def money(value):
    if value < 0:
        return f"(£{abs(value):,.0f})"
    return f"£{value:,.0f}"


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Quaerens Ltd")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 82, 164)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("12-Month P&L Forecast and Collaboration Proposal")
    r.bold = True
    r.font.size = Pt(15)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Draft management forecast for discussion with prospective legal partners")
    r.italic = True
    r.font.size = Pt(10)

    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "Quaerens Ltd operates a consumer support and assessment platform focused on identifying individuals "
        "who may have suffered financial loss arising from property, financial, travel and consumer-related issues. "
        "The purpose of this document is to provide a practical 12-month forecast showing the anticipated enquiry "
        "pipeline, likely operating costs and proposed collaboration model with a specialist UK law firm."
    )
    doc.add_paragraph(
        "This is not a filed statutory account. It is a working management forecast based on current infrastructure, "
        "existing web assets, expected marketing activity and planned growth across property and financial claim sectors."
    )

    add_heading(doc, "2. Existing Infrastructure", 1)
    infra = [
        "Quaerens.co.uk live consumer platform",
        "Specialist landing pages and topic hubs",
        "Lead capture forms and callback process",
        "CRM and client database workflow",
        "Telephone support and intake capability",
        "Organic search, social media and paid campaign routes",
        "Document collection and case screening process",
        "Multi-sector consumer education content",
    ]
    for item in infra:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "3. Target Case Categories", 1)
    add_table(
        doc,
        ["Category", "Summary", "Estimated Monthly Enquiries"],
        [
            ["Spray Foam Insulation", "Mortgage, survey, property value and removal-cost concerns", "30 - 60"],
            ["Sale & Rent Back", "Undervalue sale, lost equity, rent increases, eviction or poor advice", "10 - 20"],
            ["Equity Release", "Suitability, alternatives, interest growth, inheritance and charges", "10 - 15"],
            ["Equity Release / Investment Losses", "Released funds placed into failed property or investment schemes", "5 - 10"],
            ["Holiday Parks / Other Consumer Cases", "Existing consumer claim areas and related enquiries", "20 - 40"],
            ["Total", "Estimated enquiry range across priority sectors", "75 - 145"],
        ],
    )

    add_heading(doc, "4. 12-Month Lead Forecast", 1)
    add_table(
        doc,
        ["Metric", "Conservative", "Base Case", "Growth Case"],
        [
            ["Total enquiries", "900", "1,320", "1,740"],
            ["Qualified enquiries", "250", "375", "500"],
            ["Suitable for legal review", "150", "250", "350"],
            ["Priority property/finance cases", "90", "160", "240"],
            ["Main growth areas", "Spray foam, sale & rent back", "Equity release, property schemes", "Paid and organic expansion"],
        ],
    )

    add_heading(doc, "5. Draft P&L Forecast", 1)
    doc.add_paragraph(
        "The figures below are a planning forecast only. Revenue will depend on the final commercial agreement, "
        "regulatory position, acceptance criteria and the number of cases the legal partner is willing and able to review. "
        "Any referral, marketing, administration or support arrangement should be documented and checked for regulatory compliance."
    )

    revenue_base = 87500
    revenue_conservative = 52500
    revenue_growth = 122500
    expenses = {
        "Paid marketing and campaign testing": 18000,
        "Website, hosting, CRM and software": 4200,
        "Content, SEO and landing page development": 6000,
        "Telephone, intake and administration": 7200,
        "Professional, compliance and accountancy": 4000,
        "Design, media and document preparation": 3000,
        "General overheads and contingency": 3600,
    }
    total_expenses = sum(expenses.values())
    add_table(
        doc,
        ["P&L Line", "Conservative", "Base Case", "Growth Case"],
        [
            ["Revenue / partner service income", money(revenue_conservative), money(revenue_base), money(revenue_growth)],
            ["Cost of sales / fulfilment support", money(-6000), money(-9000), money(-12000)],
            ["Gross profit", money(revenue_conservative - 6000), money(revenue_base - 9000), money(revenue_growth - 12000)],
            ["Operating expenses", money(-total_expenses), money(-total_expenses), money(-total_expenses)],
            ["Forecast operating profit", money(revenue_conservative - 6000 - total_expenses), money(revenue_base - 9000 - total_expenses), money(revenue_growth - 12000 - total_expenses)],
        ],
    )

    doc.add_paragraph("Operating expense assumptions:")
    for k, v in expenses.items():
        doc.add_paragraph(f"{k}: {money(v)}", style="List Bullet")

    add_heading(doc, "6. Proposed Collaboration Model", 1)
    add_table(
        doc,
        ["Quaerens Responsibilities", "Law Firm Responsibilities"],
        [
            ["Lead generation and consumer awareness", "Legal merits assessment"],
            ["Initial client screening and triage", "Regulatory complaints and legal advice"],
            ["Information gathering and document collection", "Litigation or settlement strategy where appropriate"],
            ["Client onboarding and communication support", "Court proceedings and formal legal representation"],
            ["Structured case summaries for review", "Compliance, client-care and regulated legal work"],
        ],
    )

    add_heading(doc, "7. Why the Partnership Works", 1)
    doc.add_paragraph(
        "Quaerens is designed to provide a scalable front-end client acquisition and assessment process, allowing legal partners "
        "to focus on legal review, advice and progression of suitable cases. The platform is particularly suited to historic "
        "property and financial product issues where clients often need help understanding what evidence matters before a legal "
        "case can be assessed properly."
    )
    for item in [
        "Existing consumer-facing brand and live web infrastructure",
        "Priority property and finance topics already built into the site",
        "Ability to screen weak enquiries before law firm review",
        "Focus on documents, timelines and client suitability",
        "Opportunity to grow lower-competition areas such as sale and rent back and equity release-linked investments",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "8. Notes and Assumptions", 1)
    for item in [
        "Figures are indicative and should be adjusted once commercial terms are agreed.",
        "Revenue assumptions are not a guarantee of income or case acceptance.",
        "Any referral, lead generation, marketing or administration fee arrangement should be checked against SRA, FCA and data protection requirements.",
        "Forecast volumes assume continued website development, targeted content, social campaigns and active intake management.",
        "The quality of cases will depend on documentary evidence, limitation, vulnerability, financial loss and legal merits.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Quaerens Ltd - Draft 12-Month Forecast for Discussion Purposes"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
